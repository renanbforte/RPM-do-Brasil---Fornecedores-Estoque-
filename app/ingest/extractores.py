"""
Extração do conteúdo BRUTO de cada tipo de arquivo.

Filosofia: o Claude é o parser universal, então aqui a gente só precisa
despejar o conteúdo de forma fiel e completa (todas as abas, todas as linhas,
tabelas viradas em texto). Não tentamos entender colunas — isso é do LLM.
"""
from __future__ import annotations

import io
from pathlib import Path

import chardet


def _decodificar(conteudo: bytes) -> str:
    """Decodifica bytes detectando o encoding (BR costuma vir cp1252/latin-1)."""
    palpite = (chardet.detect(conteudo[:20000]) or {}).get("encoding")
    for enc in (palpite, "utf-8", "cp1252", "latin-1"):
        if not enc:
            continue
        try:
            return conteudo.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return conteudo.decode("latin-1", errors="replace")


def _extrair_txt(conteudo: bytes) -> str:
    return _decodificar(conteudo)


def _extrair_csv(conteudo: bytes) -> str:
    # Mantém o texto bruto; o separador (; ou ,) e o layout ficam para o LLM.
    return _decodificar(conteudo)


def _extrair_xlsx(conteudo: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(conteudo), read_only=True, data_only=True)
    partes: list[str] = []
    for aba in wb.sheetnames:
        ws = wb[aba]
        partes.append(f"### ABA: {aba}")
        for row in ws.iter_rows(values_only=True):
            celulas = ["" if c is None else str(c) for c in row]
            if any(c.strip() for c in celulas):
                partes.append(" | ".join(celulas))
    wb.close()
    return "\n".join(partes)


def _extrair_xls(conteudo: bytes) -> str:
    import xlrd

    wb = xlrd.open_workbook(file_contents=conteudo)
    partes: list[str] = []
    for aba in wb.sheets():
        partes.append(f"### ABA: {aba.name}")
        for r in range(aba.nrows):
            celulas = [str(aba.cell_value(r, c)) for c in range(aba.ncols)]
            if any(x.strip() for x in celulas):
                partes.append(" | ".join(celulas))
    return "\n".join(partes)


def _extrair_docx(conteudo: bytes) -> str:
    import docx

    documento = docx.Document(io.BytesIO(conteudo))
    partes: list[str] = [p.text for p in documento.paragraphs if p.text.strip()]
    for tabela in documento.tables:
        for row in tabela.rows:
            celulas = [c.text for c in row.cells]
            if any(x.strip() for x in celulas):
                partes.append(" | ".join(celulas))
    return "\n".join(partes)


_EXTRATORES = {
    ".txt": _extrair_txt,
    ".csv": _extrair_csv,
    ".xlsx": _extrair_xlsx,
    ".xls": _extrair_xls,
    ".docx": _extrair_docx,
}


def extrair_texto(nome: str, conteudo: bytes) -> str:
    """
    Extrai o texto bruto conforme a extensão do arquivo.
    Trata dupla extensão (ex: "ARQUIVO.TXT.TXT") usando o último sufixo.
    """
    ext = Path(nome).suffix.lower()
    extrator = _EXTRATORES.get(ext)
    if not extrator:
        raise ValueError(f"Extensão não suportada: {ext!r} (arquivo {nome!r})")
    return extrator(conteudo)
